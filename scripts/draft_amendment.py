#!/usr/bin/env python3
"""Генерує проєкт наказу-зміни до Наказу МОУ № 402 у форматі Word (.docx).

Структура згенерованого документа відповідає типовому проєкту НПА МОУ:

    1. Шапка: герб (placeholder), назва міністерства, тип акта.
    2. Реквізити: дата, місце, номер.
    3. Заголовок: «Про внесення змін до Положення про…»
    4. Преамбула: «Відповідно до… з метою… НАКАЗУЮ:»
    5. Основні пункти наказу-зміни (для кожного affected id):
       - insert: «Доповнити… новим пунктом такого змісту: …»
       - edit: «У пункті … слова «…» замінити словами «…»»
       - redaction: «Пункт … викласти в такій редакції: …»
       - repeal: «Пункт … виключити»
    6. Пункт про контроль виконання.
    7. Реєстрація / гриф погодження.
    8. Підпис.
    9. ДОДАТОК — порівняльна таблиця: поточна редакція || запропонована.

Використання:

    # Інтерактивно
    python3 scripts/draft_amendment.py

    # З конкретним JSON-описом
    python3 scripts/draft_amendment.py --input draft.yaml

    # Зрендерити реальний амендмент з amendments.yaml як зразок
    python3 scripts/draft_amendment.py --from-history z0329-25
    python3 scripts/draft_amendment.py --from-history z0616-24
"""

from __future__ import annotations

import argparse
import re
import sys
from datetime import date, datetime
from pathlib import Path

try:
    import yaml
except ImportError:
    sys.stderr.write("ERROR: PyYAML не встановлено. `pip install pyyaml`.\n")
    sys.exit(2)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx не встановлено. `pip install python-docx`.\n"
    )
    sys.exit(2)


REPO = Path(__file__).resolve().parent.parent
AMENDMENTS_YAML = REPO / "meta" / "amendments.yaml"
POLOZHENNIA_DIR = REPO / "polozhennia"


# ────────────────────────────────────────────────────────────────────
# Стилі документа
# ────────────────────────────────────────────────────────────────────

UA_OP_LABELS = {
    "insert": "доповнення (новий пункт)",
    "edit": "правка",
    "redaction": "нова редакція",
    "repeal": "скасування",
    "restore": "відновлення",
}

UA_OP_VERBS = {
    "insert": "доповнити",
    "edit": "змінити",
    "redaction": "викласти в новій редакції",
    "repeal": "виключити",
    "restore": "відновити дію",
}


def set_run_font(run, *, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Кириличний фолбек для шрифту
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def set_cell_border(cell, **kwargs):
    """Встановлює рамки клітинки таблиці. kwargs: top/bottom/left/right = {sz, val, color}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            attrs = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            for key, val in attrs.items():
                tag.set(qn(f"w:{key}"), str(val))
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def add_paragraph(doc, text, *, bold=False, italic=False, align=None, size=12,
                  space_before=0, space_after=6, color=None, indent_left=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


# ────────────────────────────────────────────────────────────────────
# Робота з джерелами даних
# ────────────────────────────────────────────────────────────────────

def load_amendments() -> dict:
    return yaml.safe_load(AMENDMENTS_YAML.read_text(encoding="utf-8"))


def find_amendment_by_reg(registration: str) -> dict | None:
    data = load_amendments()
    for a in data.get("amendments", []):
        if a.get("registration") == registration:
            return a
    return None


def load_chapter_for_id(target_id: str) -> tuple[dict, str] | None:
    """Шукає сторінку (frontmatter + body), яка містить пункт з даним id або
    є його батьківською главою. Повертає (frontmatter, body) або None."""
    candidates = list(POLOZHENNIA_DIR.rglob("*.md"))
    for path in candidates:
        text = path.read_text(encoding="utf-8")
        if not text.startswith("---"):
            continue
        end = text.find("\n---", 3)
        if end < 0:
            continue
        try:
            fm = yaml.safe_load(text[3:end])
        except yaml.YAMLError:
            continue
        if not isinstance(fm, dict):
            continue
        fid = fm.get("id", "")
        if fid == target_id or (fid and target_id.startswith(fid + ".")):
            body = text[end + 4:].lstrip("\n")
            return fm, body
    return None


def extract_punkt_text(body: str, punkt_id: str) -> str | None:
    """З body глави виокремлює текст конкретного пункту за маркером
    <a id="pN.M">…</a>. Текст — від цього якоря до наступного якоря або
    горизонтальної риски `---`."""
    # punkt_id типу "polozhennia.r1.gl1.p1.4" → шукаємо anchor "p1.4"
    m = re.search(r"\.p(\d[\d.]*)\.?([а-яіїєґ])?$", punkt_id)
    if not m:
        return None
    anchor = "p" + m.group(1)
    pattern = re.compile(
        rf'<a\s+id="{re.escape(anchor)}"></a>\s*\n+(.*?)(?=\n<a\s+id="|\n---\n|\Z)',
        re.DOTALL,
    )
    found = pattern.search(body)
    if not found:
        return None
    text = found.group(1).strip()
    # Прибрати маркдаунний заголовок ## N.M. з початку
    text = re.sub(r"^##\s+\d+(\.\d+)*\.?\s*\n+", "", text)
    return text


# ────────────────────────────────────────────────────────────────────
# Збірка документа
# ────────────────────────────────────────────────────────────────────

def build_document(spec: dict, out_path: Path) -> None:
    """spec має поля:
        order:       номер наказу-зміни (рядок)
        signed_at:   дата підписання (date або рядок)
        registration: zNNNN-YY реєстрація в Мін'юсті
        summary:     один-два речення
        rationale:   обґрунтування (може бути порожнім)
        affects:     список [{id, op, scope?, new_text?}]
        author:      «Міністр оборони України …»
        place:       «м. Київ»
    """
    doc = Document()
    # ── Стиль за замовчуванням ──
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # Кириличний фолбек
    rpr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rpr.append(rFonts)

    # ── Поля сторінки ──
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # ── 1. Герб + назва міністерства ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ герб України ]")
    set_run_font(run, size=10, italic=True, color=(0x80, 0x80, 0x80))

    add_paragraph(
        doc, "МІНІСТЕРСТВО ОБОРОНИ УКРАЇНИ",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=2,
    )
    add_paragraph(
        doc, "НАКАЗ",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16,
        space_before=8, space_after=14,
    )

    # ── 2. Реквізити (дата, місце, номер) у трьох колонках ──
    signed_at = spec.get("signed_at")
    if isinstance(signed_at, str):
        try:
            signed_at = datetime.fromisoformat(signed_at).date()
        except ValueError:
            pass
    date_str = (
        signed_at.strftime("%d %B %Y р.").replace("January", "січня")
        .replace("February", "лютого").replace("March", "березня")
        .replace("April", "квітня").replace("May", "травня")
        .replace("June", "червня").replace("July", "липня")
        .replace("August", "серпня").replace("September", "вересня")
        .replace("October", "жовтня").replace("November", "листопада")
        .replace("December", "грудня")
        if isinstance(signed_at, date) else (str(signed_at) if signed_at else "____________ 20__ року")
    )
    place = spec.get("place", "м. Київ")
    order_no = spec.get("order", "____")

    t = doc.add_table(rows=1, cols=3)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.autofit = True
    cells = t.rows[0].cells
    for c, txt, align in [
        (cells[0], date_str, WD_ALIGN_PARAGRAPH.LEFT),
        (cells[1], place, WD_ALIGN_PARAGRAPH.CENTER),
        (cells[2], f"№ {order_no}", WD_ALIGN_PARAGRAPH.RIGHT),
    ]:
        c.paragraphs[0].alignment = align
        r = c.paragraphs[0].add_run(txt)
        set_run_font(r, size=12)
        # Прибрати рамки
        for edge in ("top", "left", "bottom", "right"):
            set_cell_border(c, **{edge: {"val": "nil", "sz": "0"}})

    doc.add_paragraph()  # відступ

    # ── 3. Заголовок ──
    title = spec.get("title", "Про внесення змін до Положення про військово-лікарську експертизу в Збройних Силах України")
    add_paragraph(
        doc, title,
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13,
        space_before=4, space_after=20,
    )

    # ── 4. Преамбула ──
    preamble = spec.get("preamble") or (
        "Відповідно до частини тринадцятої статті 2 Закону України "
        "«Про військовий обов'язок і військову службу» та з метою "
        + (spec.get("purpose") or "вдосконалення процедури військово-лікарської експертизи")
        + ","
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(preamble + " ")
    set_run_font(r, size=12)
    r2 = p.add_run("НАКАЗУЮ:")
    set_run_font(r2, size=12, bold=True)

    # ── 5. Основні пункти наказу-зміни ──
    add_paragraph(doc, "", space_after=0)
    affects = spec.get("affects", [])
    if affects:
        intro = (
            "1. Внести до Положення про військово-лікарську експертизу "
            "в Збройних Силах України, затвердженого наказом Міністерства "
            "оборони України від 14 серпня 2008 року № 402, "
            "зареєстрованого в Міністерстві юстиції України 26 серпня 2008 року "
            "за № 1109/15800 (із змінами), такі зміни:"
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(intro)
        set_run_font(r, size=12)

        for i, aff in enumerate(affects, start=1):
            op = aff.get("op", "edit")
            target_id = aff.get("id", "")
            verb = UA_OP_VERBS.get(op, "змінити")
            # Дружній опис місця змін
            scope = aff.get("scope", "")
            location = _human_locator(target_id)
            new_text = aff.get("new_text") or "[ ТУТ_НОВА_РЕДАКЦІЯ ]"

            if op == "insert":
                clause = (
                    f"1.{i}) Доповнити {location} новим пунктом такого змісту:"
                )
            elif op == "redaction":
                clause = (
                    f"1.{i}) Викласти {location} у такій редакції:"
                )
            elif op == "repeal":
                clause = (
                    f"1.{i}) Виключити {location}."
                )
            elif op == "restore":
                clause = (
                    f"1.{i}) Відновити дію {location} у попередній редакції."
                )
            else:  # edit
                old_phrase = aff.get("old_phrase", "[слова, що замінюються]")
                new_phrase = aff.get("new_phrase", new_text)
                clause = (
                    f"1.{i}) У {location} слова «{old_phrase}» "
                    f"замінити словами «{new_phrase}»."
                )

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(clause)
            set_run_font(r, size=12)

            if op in ("insert", "redaction") and new_text:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.left_indent = Cm(1.25)
                p.paragraph_format.line_spacing = 1.5
                r = p.add_run("«" + new_text + "».")
                set_run_font(r, size=12, italic=True)

    # ── 6. Пункт контролю ──
    add_paragraph(doc, "", space_after=0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "2. Контроль за виконанням цього наказу покласти на заступника "
        "Міністра оборони України згідно з розподілом обов'язків."
    )
    set_run_font(r, size=12)

    # ── 7. Прикінцеві ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(
        "3. Цей наказ набирає чинності з дня його офіційного опублікування."
    )
    set_run_font(r, size=12)

    # ── 8. Підпис ──
    add_paragraph(doc, "", space_after=24)
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    left, right = t.rows[0].cells
    p1 = left.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p1.add_run("Міністр оборони України")
    set_run_font(r, size=12)
    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run(spec.get("author", "________________ (І. П. Прізвище)"))
    set_run_font(r, size=12)
    for c in (left, right):
        for edge in ("top", "left", "bottom", "right"):
            set_cell_border(c, **{edge: {"val": "nil", "sz": "0"}})

    # ── 9. ДОДАТОК: порівняльна таблиця ──
    doc.add_page_break()
    add_paragraph(
        doc, "ПОРІВНЯЛЬНА ТАБЛИЦЯ",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, space_after=2,
    )
    add_paragraph(
        doc, "до проєкту наказу Міністерства оборони України "
             f"«{title}»",
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, space_after=14,
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    hdr[0].width = Cm(2.5)
    hdr[1].width = Cm(7.5)
    hdr[2].width = Cm(7.5)
    for i, h in enumerate(("Пункт", "Поточна редакція", "Запропонована редакція")):
        hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = hdr[i].paragraphs[0].add_run(h)
        set_run_font(r, size=11, bold=True)
        # фон шапки
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "EAE6F5")
        tcPr.append(shd)

    for aff in affects:
        target_id = aff.get("id", "")
        op = aff.get("op", "edit")
        # Дістати поточний текст
        cur = _fetch_current_text(target_id, op)
        proposed = aff.get("new_text") or _proposed_placeholder(op)

        row = table.add_row().cells
        row[0].width = Cm(2.5)
        row[1].width = Cm(7.5)
        row[2].width = Cm(7.5)
        row[0].paragraphs[0].add_run(target_id).font.name = "Consolas"
        row[0].paragraphs[0].runs[0].font.size = Pt(9)
        # Поточна
        p = row[1].paragraphs[0]
        r = p.add_run(cur or "[пункт відсутній — буде додано]")
        set_run_font(r, size=10, italic=(cur is None))
        # Запропонована
        p = row[2].paragraphs[0]
        r = p.add_run(proposed)
        set_run_font(r, size=10, bold=(op in ("insert", "redaction")))
        if op == "repeal":
            # Закреслюємо
            r.font.strike = True
            r.text = cur or "[зміст пункту]"

    # Підпис до таблиці
    add_paragraph(doc, "", space_after=10)
    add_paragraph(
        doc, f"Розробник: {spec.get('drafter', '_______________')}",
        align=WD_ALIGN_PARAGRAPH.LEFT, size=11,
    )

    # ── ДОДАТОК 2: метадані для трекінгу ──
    doc.add_page_break()
    add_paragraph(
        doc, "ДОВІДКОВІ ВІДОМОСТІ",
        bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, space_after=14,
    )

    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    info_table.autofit = True

    info_rows = [
        ("Реєстрація в Мін'юсті (планова)", spec.get("registration") or "z____-__"),
        ("Джерело-наказ", "МОУ № 402 від 14.08.2008 (z1109-08)"),
        ("Кількість зачеплених пунктів", str(len(affects))),
        ("Типи операцій", ", ".join(sorted({UA_OP_LABELS.get(a.get("op", "edit"), "?") for a in affects}))),
        ("Резюме", spec.get("summary") or "—"),
        ("Обґрунтування", spec.get("rationale") or "[ заповнити у пояснювальній записці ]"),
    ]
    for label, value in info_rows:
        row = info_table.add_row().cells
        r1 = row[0].paragraphs[0].add_run(label)
        set_run_font(r1, size=11, bold=True)
        r2 = row[1].paragraphs[0].add_run(value)
        set_run_font(r2, size=11)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _human_locator(target_id: str) -> str:
    """polozhennia.r1.gl1.p1.4 → 'пункт 1.4 глави 1 розділу I'"""
    m = re.match(
        r"polozhennia\.r(\d+)(?:\.gl(\d+))?(?:\.p([\d.]+))?(?:\.([а-яіїєґ]))?",
        target_id,
    )
    if not m:
        m2 = re.match(r"nakaz\.p(\d+)", target_id)
        if m2:
            return f"пункт {m2.group(1)} наказу"
        m3 = re.match(r"dodatok\.(\d+)(?:\.stattia\.(\d+))?", target_id)
        if m3:
            d = m3.group(1)
            if m3.group(2):
                return f"статтю {m3.group(2)} Додатка {d}"
            return f"Додаток {d}"
        return target_id
    roman = ["", "I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]
    rozdil = roman[int(m.group(1))] if int(m.group(1)) < len(roman) else m.group(1)
    parts = [f"розділ {rozdil}"]
    if m.group(2):
        parts.append(f"глава {m.group(2)}")
    if m.group(3):
        parts.append(f"пункт {m.group(3)}")
    if m.group(4):
        parts.append(f"підпункт «{m.group(4)}»")
    # розвернути в зворотньому порядку для природного звучання
    return " ".join(parts[::-1])


def _fetch_current_text(target_id: str, op: str) -> str | None:
    if op == "insert":
        return None
    found = load_chapter_for_id(target_id)
    if not found:
        return None
    fm, body = found
    if fm.get("id") == target_id:
        return body[:600].strip() + ("…" if len(body) > 600 else "")
    text = extract_punkt_text(body, target_id)
    if text:
        # Прибрати markdown-розмітку для читабельності в .docx
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r">\s*\*\{[^}]*\}\*", "", text)  # маркери амендментів
        text = re.sub(r"\n\s*\n+", "\n\n", text)
        return text[:800].strip() + ("…" if len(text) > 800 else "")
    return None


def _proposed_placeholder(op: str) -> str:
    return {
        "insert":   "[ ТЕКСТ_НОВОГО_ПУНКТУ ]",
        "edit":     "[ ТЕКСТ_ПІСЛЯ_ПРАВКИ ]",
        "redaction": "[ ТЕКСТ_НОВОЇ_РЕДАКЦІЇ ]",
        "repeal":   "(пункт виключається)",
        "restore":  "[ ВІДНОВЛЮЄТЬСЯ_ПОПЕРЕДНЯ_РЕДАКЦІЯ ]",
    }.get(op, "[ ЗМІСТ_ПРАВКИ ]")


# ────────────────────────────────────────────────────────────────────
# CLI
# ────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", help="YAML-файл зі spec'ом (див. --help-spec)")
    parser.add_argument("--from-history", help="Реєстрація амендменту (zNNNN-YY) з amendments.yaml")
    parser.add_argument("--out", help="Output .docx", default=None)
    args = parser.parse_args()

    if args.from_history:
        a = find_amendment_by_reg(args.from_history)
        if a is None:
            sys.stderr.write(f"Не знайдено амендмент {args.from_history} у amendments.yaml\n")
            return 1
        spec = {
            "order": a["order"],
            "signed_at": a.get("signed_at"),
            "registration": a.get("registration"),
            "summary": a.get("summary"),
            "affects": a.get("affects", []),
            "title": _title_from_summary(a.get("summary", "")),
        }
        out = args.out or REPO / "examples" / "drafts" / f"draft-{a['order']}-{args.from_history}.docx"
    elif args.input:
        spec = yaml.safe_load(Path(args.input).read_text(encoding="utf-8"))
        out = args.out or REPO / "dist" / f"draft-{spec.get('order', 'unknown')}.docx"
    else:
        sys.stderr.write(
            "Вкажіть --input <yaml> або --from-history <zNNNN-YY>.\n"
            "Приклад: python3 scripts/draft_amendment.py --from-history z0329-25\n"
        )
        return 1

    out = Path(out)
    build_document(spec, out)
    print(f"Згенеровано: {out.relative_to(REPO) if out.is_relative_to(REPO) else out}")
    return 0


def _title_from_summary(summary: str) -> str:
    """Витягує заголовок «Про внесення змін…» з summary або повертає типовий."""
    return "Про внесення змін до Положення про військово-лікарську експертизу в Збройних Силах України"


if __name__ == "__main__":
    sys.exit(main())
